from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
REPORT_MD = DOCS / "CS599_大作业报告.md"
REPORT_DOCX = DOCS / "CS599_大作业报告.docx"
REPORT_PDF = DOCS / "CS599_大作业报告.pdf"


def set_run_font(run, size=None, bold=None, name="宋体", color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=12, name="宋体", bold=None, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, name=name, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bookmark(paragraph, name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str, size=12, bold=False, name="宋体", color="111827"):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    r_pr.append(fonts)

    for tag in ("w:sz", "w:szCs"):
        size_node = OxmlElement(tag)
        size_node.set(qn("w:val"), str(int(size * 2)))
        r_pr.append(size_node)

    color_node = OxmlElement("w:color")
    color_node.set(qn("w:val"), color)
    r_pr.append(color_node)

    if bold:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)

    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(r_pr)
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_toc_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for node in (begin, instr, separate):
        paragraph.add_run()._r.append(node)
    run = paragraph.add_run("请在 Word 中右键此处，选择“更新域/更新目录”生成目录。")
    set_run_font(run, size=12, name="宋体", color="6B7280")
    paragraph.add_run()._r.append(end)


def enable_update_fields_on_open(document: Document):
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_document(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(4)

    h1 = styles["Heading 1"]
    h1.font.name = "黑体"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("111827")
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)

    h2 = styles["Heading 2"]
    h2.font.name = "黑体"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string("1F2937")
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(8)

    h3 = styles["Heading 3"]
    h3.font.name = "黑体"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    h3.font.size = Pt(12.5)
    h3.font.bold = True


def add_page_number(section):
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=10, name="Times New Roman")


def read_pdf_page_map():
    if PdfReader is None or not REPORT_PDF.exists():
        return {}
    try:
        reader = PdfReader(str(REPORT_PDF))
        page_map = {}

        def walk(items):
            for item in items:
                if isinstance(item, list):
                    walk(item)
                    continue
                title = str(getattr(item, "title", "")).strip()
                if title:
                    page_map[title] = reader.get_destination_page_number(item) + 1

        walk(reader.outline)
        return page_map
    except Exception:
        return {}


def collect_toc_items(markdown: str):
    items = []
    heading_index = 0
    for line in markdown.splitlines():
        stripped = line.strip()
        if stripped.startswith("## ") and stripped != "## 目录":
            heading_index += 1
            items.append((stripped[3:], 1, f"bm_h_{heading_index}"))
        elif stripped.startswith("### "):
            heading_index += 1
            items.append((stripped[4:], 2, f"bm_h_{heading_index}"))
    return items


def add_cover(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(52)
    run = p.add_run("武汉理工大学")
    set_run_font(run, size=24, bold=True, name="黑体", color="111827")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    run = p.add_run("CS599 期末大作业报告")
    set_run_font(run, size=22, bold=True, name="黑体", color="111827")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(34)
    run = p.add_run("法律责任初步分析 Agent")
    set_run_font(run, size=20, bold=True, name="黑体", color="1D4ED8")

    table = document.add_table(rows=8, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    data = [
        ("课程名称", "企业级应用软件设计与开发"),
        ("项目名称", "法律责任初步分析 Agent"),
        ("方向", "方向一：Agentic AI 原生开发"),
        ("学号", "2025302937"),
        ("姓名", "宋怡康"),
        ("专业", "计算机技术"),
        ("指导教师", "威欣"),
        ("提交日期", "2026 年 6 月 22 日"),
    ]
    for idx, (row, (key, value)) in enumerate(zip(table.rows, data)):
        row.cells[0].width = Cm(4.4)
        row.cells[1].width = Cm(9.8)
        row.cells[0].text = key
        row.cells[1].text = value
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=130, bottom=130, left=160, right=160)
            set_cell_shading(cell, "EEF2F7" if c_idx == 0 else "FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_font(paragraph, size=12, bold=(c_idx == 0), color="111827")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    run = p.add_run("按武汉理工大学硕士学位论文参考格式排版")
    set_run_font(run, size=11, name="宋体", color="4B5563")

    document.add_page_break()


def add_catalog(document: Document, markdown: str):
    p = document.add_paragraph("目录")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(24)
    set_paragraph_font(p, size=16, name="黑体", bold=True)
    add_bookmark(p, "toc", 1)

    field_paragraph = document.add_paragraph()
    field_paragraph.paragraph_format.line_spacing = 1.15
    field_paragraph.paragraph_format.space_after = Pt(8)
    add_toc_field(field_paragraph)

    document.add_page_break()


def parse_inline(paragraph, text, size=12):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size=10.5, name="Consolas", color="1D4ED8")
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=size, name="宋体")


def add_table(document: Document, rows):
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, "DBEAFE")
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_font(paragraph, size=10.2, bold=(r_idx == 0), color="111827")
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(document: Document, code_lines):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    p = cell.paragraphs[0]
    run = p.add_run("\n".join(code_lines))
    set_run_font(run, size=9.5, name="Consolas", color="111827")
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_body(document: Document, markdown: str):
    lines = markdown.splitlines()
    i = 0
    in_code = False
    code_lines = []
    skip_initial_title = True
    heading_index = 0
    bookmark_id = 10

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                in_code = False
                code_lines = []
            else:
                in_code = True
                code_lines = []
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if skip_initial_title and stripped.startswith("# "):
            skip_initial_title = False
            i += 1
            continue

        if stripped == "## 目录":
            while i < len(lines) and not lines[i].startswith("## 摘要"):
                i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|") and "---" in lines[i + 1]:
            table_lines = [stripped]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            add_table(document, rows)
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if image_match:
            image_path = ROOT / image_match.group(2)
            if image_path.exists():
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(image_path), width=Cm(14.8))
                cap = document.add_paragraph(image_match.group(1))
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_font(cap, size=10.5, color="4B5563")
            i += 1
            continue

        if stripped.startswith("## "):
            heading_index += 1
            p = document.add_paragraph(stripped[3:], style="Heading 1")
            add_bookmark(p, f"bm_h_{heading_index}", bookmark_id)
            bookmark_id += 1
            if stripped == "## 摘要":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(50)
                p.paragraph_format.space_after = Pt(22)
            set_paragraph_font(p, size=16, name="黑体", bold=True, color="111827")
        elif stripped.startswith("### "):
            heading_index += 1
            p = document.add_paragraph(stripped[4:], style="Heading 2")
            add_bookmark(p, f"bm_h_{heading_index}", bookmark_id)
            bookmark_id += 1
            set_paragraph_font(p, size=14, name="黑体", bold=True, color="1F2937")
        elif stripped.startswith("- "):
            p = document.add_paragraph(style="List Paragraph")
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            parse_inline(p, "• " + stripped[2:], size=11.5)
        elif re.match(r"^\d+\.\s+", stripped):
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            parse_inline(p, stripped)
        else:
            p = document.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(5)
            parse_inline(p, stripped)
        i += 1


def main():
    document = Document()
    configure_document(document)
    enable_update_fields_on_open(document)
    add_page_number(document.sections[0])
    add_cover(document)
    markdown = REPORT_MD.read_text(encoding="utf-8")
    add_catalog(document, markdown)
    add_body(document, markdown)
    document.save(str(REPORT_DOCX))
    print(REPORT_DOCX)


if __name__ == "__main__":
    main()
